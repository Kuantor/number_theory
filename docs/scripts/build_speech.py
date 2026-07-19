"""Build the YouTube-friendly presentation speech (issue #20) as a .docx."""
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches

OUT = sys.argv[1]

NAVY = RGBColor(0x1E, 0x2A, 0x3A)
GOLD = RGBColor(0xA8, 0x7C, 0x10)
BLUE = RGBColor(0x2E, 0x6B, 0xA0)
GRAY = RGBColor(0x70, 0x70, 0x70)
BLACK = RGBColor(0x22, 0x22, 0x22)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(12)
normal.font.color.rgb = BLACK
doc.sections[0].left_margin = Inches(1.0)
doc.sections[0].right_margin = Inches(1.0)
doc.sections[0].top_margin = Inches(0.9)
doc.sections[0].bottom_margin = Inches(0.9)


def add(text="", size=12, color=BLACK, bold=False, italic=False, name="Calibri",
        after=8, before=0, align=None, ls=1.18):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing = ls
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.name = name
    r.font.color.rgb = color
    return p


def slide_head(n, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("SLIDE %d" % n)
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.name = "Consolas"
    r.font.color.rgb = GOLD
    r2 = p.add_run("   " + title)
    r2.font.size = Pt(15)
    r2.font.bold = True
    r2.font.name = "Calibri"
    r2.font.color.rgb = NAVY


def cue(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run("▸ " + text)
    r.font.size = Pt(10.5)
    r.font.italic = True
    r.font.color.rgb = GRAY


# ---------------------------------------------------------------- header
add("number_theory", size=26, color=NAVY, bold=True, name="Consolas", after=2)
add("Presentation speech / spoken script", size=14, color=BLUE, italic=True, after=10)
add("Approx. runtime: 9–11 minutes at a relaxed pace.  Tone: upbeat, conversational, "
    "a little nerdy.  The ▸ lines are delivery cues, not spoken.  Jokes are baked in — "
    "keep them, or drop any that don't fit your style.", size=11, color=GRAY, italic=True, after=4)
add("_" * 60, size=11, color=RGBColor(0xC8, 0xD0, 0xD8), after=6)

# ---------------------------------------------------------------- intro
slide_head(1, "Title")
cue("Warm, welcoming. Smile. Let the title slide breathe for a second.")
add("Hey everyone, welcome! Today I want to show you a little Python library I built, "
    "called number_theory. Right now it does two things — greatest common divisor and "
    "least common multiple — and it does them with a twist you probably won't expect.")
add("And before your eyes glaze over at the words “number theory” — stick with me. "
    "This is the branch of math behind your online banking, your Wi-Fi password, and every "
    "secure message you’ve ever sent. Also, there will be jokes. Some of them will even "
    "be good. Let’s dive in.")

slide_head(2, "A tiny library with a big twist")
add("So, what is number_theory? At its heart, it’s a tiny, dependency-free Python "
    "library. No giant frameworks, no four hundred packages downloading when you install it "
    "— zero runtime dependencies. You import it, you call a function, you get an answer.")
add("Version 0.1.0 is out in the world today, and it ships five functions. But here’s the "
    "twist that makes this more than a homework assignment…")

slide_head(3, "Ship the math, not the source")
add("Normally, when you pip install someone’s Python package, you get their source code. "
    "All of it. Right there in plain text — ready to read, copy, or quietly judge their "
    "variable names.")
add("number_theory does something different: it ships the math as a compiled binary. Users "
    "get working functions; the actual implementation is compiled machine code. Why would "
    "you do that? Two big reasons — protecting your intellectual property, and shipping an "
    "internal tool to someone without handing over the recipe.")
cue("Beat, then the punchline.")
add("We compiled the source so thoroughly that even we can’t read it anymore… I’m "
    "kidding. It’s in Git. Calm down.")

slide_head(4, "What's shipping — and what's coming")
add("Here’s the honest status. Shipping today: gcd, lcm, extended_gcd, and the "
    "“many” versions that take a whole list of numbers at once. On the roadmap: "
    "primality testing, a prime sieve, factorization, Euler’s totient, and modular "
    "arithmetic.")
add("Think of today as the appetizer, and the roadmap as the… larger appetizer. It’s a "
    "growing menu.")

slide_head(5, "The MVP: GCD & LCM, the Euclidean way")
add("Let’s talk about the actual math, because it’s genuinely beautiful. The greatest "
    "common divisor — the biggest number that divides two others — is computed with "
    "Euclid’s algorithm. You replace the pair (a, b) with (b, a mod b), and repeat until "
    "the second number hits zero. That’s the whole thing.")
add("Here’s the wild part: this algorithm is about two thousand three hundred years old. "
    "Euclid wrote it down around 300 BC. Your brand-new CPU, running Python 3.14, is "
    "executing an algorithm older than the Roman Empire — and it’s still optimal.")
cue("Fact — say it like you’re sharing a secret.")
add("Bonus fact: the absolute worst case for Euclid’s algorithm — the inputs that make it "
    "work the hardest — are consecutive Fibonacci numbers. Math is basically one big "
    "cinematic universe; everything connects. And extended_gcd hands you back the "
    "coefficients where a-times-x plus b-times-y equals the gcd — which turns out to be the "
    "seed of RSA encryption. Hold that thought.")

slide_head(6, "The math still to come")
add("The roadmap is where it gets spicy. Primality testing: is this number prime? "
    "Factorization: break a number into its prime building blocks — sixty is two times two "
    "times three times five. A prime sieve to generate primes, named after Eratosthenes, a "
    "Greek librarian who also measured the circumference of the Earth using a stick and some "
    "shadows. Overachiever.")
add("Euler’s totient counts how many numbers below n share no common factor with it. And "
    "modular arithmetic is just clock math: on a twelve-hour clock, ten o’clock plus five "
    "hours is three, not fifteen.")
cue("Crowd-pleaser fact.")
add("Here’s one to drop at parties: there are infinitely many prime numbers, and Euclid "
    "proved it around 300 BC with a proof so clean it still shows up in textbooks today. "
    "Primes are basically the introverts of math — they refuse to share factors with anyone.")

slide_head(7, "What is a Python wheel?")
add("Quick technical detour number one, because I promised you’d learn something. What is "
    "a Python wheel? A wheel — file extension dot-w-h-l — is Python’s built package "
    "format. And I want to fully demystify it: a wheel is just a ZIP file wearing a name tag. "
    "That’s it.")
add("“Built” means it’s ready to install — pip just unzips it into place, no "
    "compiling on your machine. And that funny-looking filename actually tells you "
    "everything: the package name, the version, cp314 meaning CPython 3.14, and win_amd64 "
    "meaning 64-bit Windows.")
add("For us, wheels are perfect, because our compiled binary just rides along inside the zip. "
    "So yes — we reinvented the wheel. This time, literally.")

slide_head(8, "Cython — and why compiling is faster")
add("Detour number two: Cython, and why compiling makes things faster. Normal Python is "
    "interpreted — the computer reads your code line by line, at runtime, every single time. "
    "Picture a very polite butler double-checking every instruction before doing it.")
add("Cython compiles Python-ish code down to C, and then to native machine code that runs "
    "straight on the CPU — that’s the butler who already knows what you want and just does "
    "it. Add static C types and tight loops can run ten to a hundred times faster.")
cue("Honesty beat — keeps you credible.")
add("Now, full honesty: in number_theory we mostly use Cython to hide the source. We kept "
    "Python’s arbitrary-precision integers, so today’s speedups are modest. But the fast "
    "lane is paved and ready for when we add the heavy numeric stuff.")

slide_head(9, "How we ship without the source")
add("So how do we actually ship without the source? The algorithms live in a file called "
    "_core.pyx — that’s the Cython. We compile it into a dot-p-y-d file, which is machine "
    "code. Then, when we build the wheel, we include only that compiled binary and a couple "
    "of thin wrapper files — and we deliberately leave out the Cython source and the "
    "generated C.")
add("And we didn’t just trust that. We unzipped the finished wheel and checked. No "
    "algorithm source in there — just the binary. It’s the software equivalent of mailing "
    "someone a working espresso machine instead of the blueprints.")

slide_head(10, "Getting it: installation")
add("Okay — how do you actually use this thing? The released wheel targets Windows 64-bit and "
    "Python 3.14, and you just pip install it straight from the GitHub release. On a "
    "different operating system or Python version, you can build it yourself from source, as "
    "long as you have a C compiler.")
add("And do yourself a favor: install into a virtual environment, so you don’t sprinkle "
    "packages all over your system Python like glitter that never comes out of the carpet.")

slide_head(11, "Using it: command line & Python")
add("There are two ways in. From the command line: “number-theory gcd 12 18 24” gives "
    "you 6. “lcm 4 6 8” gives you 24. Clean and quick.")
add("Or from Python: import number_theory as nt, and call nt.gcd, nt.lcm, and friends. Same "
    "math, whichever door you walk through.")

slide_head(12, "Gotchas we hit — so you won't")
add("Let me save you three headaches I already suffered through. One: “number-theory is "
    "not recognized.” That’s not broken — Python’s Scripts folder just isn’t on "
    "your PATH. Use python-dash-m number_theory, or add the folder.")
add("Two: “not a supported wheel on this platform.” The release wheel is Windows-and-"
    "Python-3.14 only; on anything else, build from source.")
cue("Land the last one lightly — it’s the funniest.")
add("And three, my personal favorite: ModuleNotFoundError for number_theory dot _core. That "
    "happens when you run Python from inside the project folder, so the source folder shadows "
    "the installed binary. The fix is basically “walk away from the repo and try "
    "again” — which, honestly, is good advice for a lot of things in life.")

slide_head(13, "Release notes — v0.1.0")
add("Release notes for version 0.1.0, nicknamed the “GCD and LCM MVP.” First release: "
    "GCD and LCM via Euclid, as a compiled, source-less binary. Inside you get gcd, lcm, "
    "extended_gcd, and the “many” variants — all arbitrary precision, so they’ll "
    "happily chew on thousand-digit numbers.")
add("Only the compiled binary ships. And it closes out four project issues, because we like a "
    "tidy changelog.")

slide_head(14, "Built green: CI & automation")
add("A quick brag about quality, because it matters. Every push and every pull request kicks "
    "off GitHub Actions, which builds the wheel and runs the tests against the actual "
    "compiled module — not the source, the real shipped binary. And it does that across three "
    "operating systems times four Python versions. Twelve combinations. All green.")
add("There’s also a workflow that automatically builds and attaches wheels for every "
    "platform whenever we cut a release. So the robots do the boring part — which is exactly "
    "what robots are for.")

slide_head(15, "What's next")
add("What’s next? More math: primality, the sieve, factorization, Euler’s totient, and "
    "modular arithmetic — and that last one is the actual engine behind RSA encryption, so "
    "we’re building toward the stuff that secures the internet.")
add("Beyond the math: portable wheels for every operating system, and maybe a PyPI release so "
    "you can just pip install number_theory with no fuss. Plus an even tighter binary. "
    "It’s all tracked as issues on the roadmap board — nothing up my sleeve.")

slide_head(16, "Closing")
cue("Slow down. Warm landing.")
add("And that’s number_theory. Two functions today, a whole constellation of them on the "
    "way. If you want to poke at it, it’s on GitHub at github-dot-com-slash-Kuantor-slash-"
    "number_theory.")
add("Remember: gcd and lcm are shipping now, and the rest of the primes are on their way — "
    "they’re just being introverts about it. Thanks so much for watching. Now go compute "
    "something.")

doc.save(OUT)
print("saved", OUT)
